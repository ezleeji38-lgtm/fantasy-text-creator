"""판타지 소설 작업 현황 대시보드.

`novelcraft dashboard` 로 실행한다. projects/ 폴더를 스캔해서
작품·챕터·Bible 완성도를 로컬 웹 UI로 보여주고,
챕터 집필·리뷰·승인·Bible 편집을 브라우저에서 바로 수행한다.
"""
from __future__ import annotations

import asyncio
import json
import re
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any

from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, StreamingResponse

from .config import PROJECTS_DIR
from .jobs import JobEvent, manager as job_manager

ALLOWED_BIBLE_FILES = {
    "bible/world.md",
    "bible/timeline.md",
    "bible/glossary.md",
    "bible/themes.md",
    "outline/synopsis.md",
    "outline/arcs.md",
    "outline/chapters.md",
    "behavior.md",
}

TEMPLATES_DIR = Path(__file__).resolve().parent / "templates"
CHAPTER_HEADER_RE = re.compile(r"^##\s+(ch\d+)\b(.*)$", re.MULTILINE)
CHAPTER_FILE_RE = re.compile(r"^(ch\d+)")


def _mtime(path: Path) -> float:
    try:
        return path.stat().st_mtime
    except OSError:
        return 0.0


def _read_text(path: Path) -> str:
    if not path.exists():
        return ""
    try:
        return path.read_text(encoding="utf-8")
    except OSError:
        return ""


def _parse_chapter_plan(chapters_md: str) -> list[dict]:
    """outline/chapters.md 의 `## chNN 제목` 블록 목록."""
    if not chapters_md:
        return []
    matches = list(CHAPTER_HEADER_RE.finditer(chapters_md))
    chapters: list[dict] = []
    for i, m in enumerate(matches):
        cid = m.group(1)
        title_line = m.group(2).strip().lstrip("—–-: ").strip()
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(chapters_md)
        body = chapters_md[start:end].strip()
        chapters.append({
            "id": cid,
            "title": title_line or cid,
            "beat": body[:280],
        })
    return chapters


def _bible_report(project_path: Path) -> dict:
    checks = {
        "world": project_path / "bible/world.md",
        "characters": project_path / "bible/characters",
        "timeline": project_path / "bible/timeline.md",
        "glossary": project_path / "bible/glossary.md",
        "themes": project_path / "bible/themes.md",
        "synopsis": project_path / "outline/synopsis.md",
        "arcs": project_path / "outline/arcs.md",
        "chapters": project_path / "outline/chapters.md",
        "behavior": project_path / "behavior.md",
    }
    items = []
    ready = 0
    for key, path in checks.items():
        if path.is_dir():
            count = len(list(path.glob("*.md")))
            ok = count > 0
            items.append({"key": key, "ok": ok, "size": count, "kind": "dir"})
        else:
            size = len(_read_text(path))
            ok = size >= 200
            items.append({"key": key, "ok": ok, "size": size, "kind": "file"})
        if ok:
            ready += 1
    return {"items": items, "ready": ready, "total": len(checks)}


def _scan_chapters(project_path: Path, planned: list[dict]) -> list[dict]:
    drafts_dir = project_path / "drafts"
    revisions_dir = project_path / "revisions"
    final_dir = project_path / "final"

    seen: dict[str, dict] = {c["id"]: c for c in planned}
    for d in (drafts_dir, revisions_dir, final_dir):
        if not d.exists():
            continue
        for f in d.glob("*.md"):
            m = CHAPTER_FILE_RE.match(f.stem)
            if m:
                cid = m.group(1)
                seen.setdefault(cid, {"id": cid, "title": cid, "beat": ""})

    def sort_key(item: dict) -> int:
        try:
            return int(item["id"][2:])
        except (ValueError, KeyError):
            return 9999

    ordered = sorted(seen.values(), key=sort_key)

    result: list[dict] = []
    for ch in ordered:
        cid = ch["id"]
        draft = drafts_dir / f"{cid}_draft.md"
        final_path = final_dir / f"{cid}.md"
        rev_files = sorted(revisions_dir.glob(f"{cid}_r*.md")) if revisions_dir.exists() else []

        final_text = _read_text(final_path) if final_path.exists() else ""
        draft_text = _read_text(draft) if draft.exists() else ""

        if final_text:
            status = "final"
            words = len(final_text)
            touched = _mtime(final_path)
        elif rev_files:
            status = "revised"
            latest = max(rev_files, key=_mtime)
            words = len(_read_text(latest))
            touched = _mtime(latest)
        elif draft_text:
            status = "draft"
            words = len(draft_text)
            touched = _mtime(draft)
        else:
            status = "planned"
            words = 0
            touched = 0.0

        result.append({
            "id": cid,
            "title": ch.get("title") or cid,
            "beat": ch.get("beat", ""),
            "status": status,
            "words": words,
            "touched": touched,
        })
    return result


def _scan_project(project_path: Path) -> dict | None:
    meta_file = project_path / "project.json"
    if not meta_file.exists():
        return None
    try:
        meta = json.loads(meta_file.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        meta = {"name": project_path.name}

    planned = _parse_chapter_plan(_read_text(project_path / "outline/chapters.md"))
    chapters = _scan_chapters(project_path, planned)

    total_words = sum(c["words"] for c in chapters)
    target = meta.get("target_length") or 100_000
    progress = min(1.0, total_words / target) if target else 0.0

    status_counts = {"final": 0, "revised": 0, "draft": 0, "planned": 0}
    for c in chapters:
        status_counts[c["status"]] = status_counts.get(c["status"], 0) + 1

    last_touched = max((c["touched"] for c in chapters), default=_mtime(meta_file))
    last_iso = datetime.fromtimestamp(last_touched).isoformat() if last_touched else None

    images = _scan_images(project_path)

    return {
        "name": meta.get("name", project_path.name),
        "genre": meta.get("genre", "판타지"),
        "target_length": target,
        "model": meta.get("model", ""),
        "total_words": total_words,
        "progress": progress,
        "chapters": chapters,
        "status_counts": status_counts,
        "bible": _bible_report(project_path),
        "images": images,
        "last_touched": last_touched,
        "last_touched_iso": last_iso,
    }


IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif"}


def _scan_images(project_path: Path) -> list[dict]:
    img_dir = project_path / "images"
    if not img_dir.exists():
        return []
    images: list[dict] = []
    for f in sorted(img_dir.iterdir()):
        if f.suffix.lower() in IMAGE_EXTS and f.is_file():
            label = f.stem.replace("_", " ")
            if label and label[0].isdigit() and " " in label:
                label = label.split(" ", 1)[1]
            images.append({
                "filename": f.name,
                "label": label.title(),
                "size_kb": round(f.stat().st_size / 1024),
            })
    return images


def scan_all_projects() -> list[dict]:
    if not PROJECTS_DIR.exists():
        return []
    out: list[dict] = []
    for child in sorted(PROJECTS_DIR.iterdir()):
        if not child.is_dir():
            continue
        data = _scan_project(child)
        if data:
            out.append(data)
    out.sort(key=lambda p: p["last_touched"], reverse=True)
    return out


def _project_path(name: str) -> Path:
    path = PROJECTS_DIR / name
    if not path.exists() or not (path / "project.json").exists():
        raise HTTPException(status_code=404, detail=f"project '{name}' not found")
    return path


def _safe_bible_relpath(relpath: str) -> str:
    """허용 파일이거나 bible/characters/ 하위 .md만 통과."""
    if relpath in ALLOWED_BIBLE_FILES:
        return relpath
    if relpath.startswith("bible/characters/") and relpath.endswith(".md"):
        if ".." in relpath or "\\" in relpath:
            raise HTTPException(status_code=400, detail="invalid path")
        return relpath
    raise HTTPException(status_code=400, detail=f"not an editable bible file: {relpath}")


def create_app() -> FastAPI:
    app = FastAPI(title="novelcraft dashboard")

    # ─────────── 페이지 & 프로젝트 스캔 ───────────
    @app.get("/", response_class=HTMLResponse)
    def index() -> HTMLResponse:
        html = (TEMPLATES_DIR / "dashboard.html").read_text(encoding="utf-8")
        return HTMLResponse(html)

    @app.get("/api/projects")
    def api_projects() -> JSONResponse:
        return JSONResponse({
            "generated_at": datetime.now().isoformat(timespec="seconds"),
            "projects": scan_all_projects(),
        })

    # ─────────── 챕터 리소스 ───────────
    @app.get("/api/projects/{name}/chapters/{chapter_id}")
    def api_chapter_detail(name: str, chapter_id: str) -> JSONResponse:
        path = _project_path(name)
        draft = path / "drafts" / f"{chapter_id}_draft.md"
        final = path / "final" / f"{chapter_id}.md"
        review = path / "drafts" / f"{chapter_id}_review.json"
        summary = path / "memory/chapter_summaries" / f"{chapter_id}_summary.md"

        planned = _parse_chapter_plan(_read_text(path / "outline/chapters.md"))
        beat = next((c["beat"] for c in planned if c["id"] == chapter_id), "")

        review_data: Any = None
        if review.exists():
            try:
                review_data = json.loads(review.read_text(encoding="utf-8"))
            except json.JSONDecodeError:
                review_data = {"parse_error": True}

        return JSONResponse({
            "chapter_id": chapter_id,
            "beat": beat,
            "draft": _read_text(draft) if draft.exists() else None,
            "final": _read_text(final) if final.exists() else None,
            "review": review_data,
            "summary": _read_text(summary) if summary.exists() else None,
            "has_draft": draft.exists(),
            "has_final": final.exists(),
        })

    @app.post("/api/projects/{name}/chapters/{chapter_id}/approve")
    def api_chapter_approve(name: str, chapter_id: str) -> JSONResponse:
        """draft를 final로 복사(승격)."""
        path = _project_path(name)
        draft = path / "drafts" / f"{chapter_id}_draft.md"
        if not draft.exists():
            raise HTTPException(status_code=404, detail="draft not found")
        final_dir = path / "final"
        final_dir.mkdir(parents=True, exist_ok=True)
        final = final_dir / f"{chapter_id}.md"
        shutil.copyfile(draft, final)
        return JSONResponse({"ok": True, "final_path": str(final)})

    # ─────────── 이미지 서빙 ───────────
    @app.get("/api/projects/{name}/images")
    def api_images_list(name: str) -> JSONResponse:
        path = _project_path(name)
        return JSONResponse({"images": _scan_images(path)})

    @app.get("/api/projects/{name}/images/{filename}")
    def api_image_file(name: str, filename: str) -> FileResponse:
        path = _project_path(name)
        if ".." in filename or "/" in filename or "\\" in filename:
            raise HTTPException(status_code=400, detail="invalid filename")
        img = path / "images" / filename
        if not img.exists() or img.suffix.lower() not in IMAGE_EXTS:
            raise HTTPException(status_code=404, detail="image not found")
        media = {
            ".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
            ".webp": "image/webp", ".gif": "image/gif",
        }.get(img.suffix.lower(), "application/octet-stream")
        return FileResponse(img, media_type=media)

    # ─────────── Bible 읽기/편집 ───────────
    @app.get("/api/projects/{name}/bible")
    def api_bible_read(name: str, file: str) -> JSONResponse:
        path = _project_path(name)
        rel = _safe_bible_relpath(file)
        target = path / rel
        if not target.exists():
            return JSONResponse({"file": rel, "text": "", "exists": False})
        return JSONResponse({
            "file": rel,
            "text": target.read_text(encoding="utf-8"),
            "exists": True,
        })

    @app.put("/api/projects/{name}/bible")
    async def api_bible_write(name: str, request: Request) -> JSONResponse:
        path = _project_path(name)
        body = await request.json()
        rel = _safe_bible_relpath(body.get("file", ""))
        text = body.get("text", "")
        if not isinstance(text, str):
            raise HTTPException(status_code=400, detail="text must be string")
        target = path / rel
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_text(text, encoding="utf-8")
        return JSONResponse({"ok": True, "file": rel, "bytes": len(text.encode("utf-8"))})

    @app.get("/api/projects/{name}/bible/list")
    def api_bible_list(name: str) -> JSONResponse:
        path = _project_path(name)
        files = sorted(ALLOWED_BIBLE_FILES)
        char_dir = path / "bible/characters"
        chars: list[str] = []
        if char_dir.exists():
            for f in sorted(char_dir.glob("*.md")):
                chars.append(f"bible/characters/{f.name}")
        return JSONResponse({"common": files, "characters": chars})

    # ─────────── DOCX 다운로드 ───────────
    @app.get("/api/projects/{name}/export/docx")
    def api_export_docx(name: str):
        """전체 챕터를 하나의 DOCX로 컴파일해 다운로드."""
        from docx import Document as DocxDocument
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import tempfile

        path = _project_path(name)
        meta_file = path / "project.json"
        meta = json.loads(meta_file.read_text(encoding="utf-8"))

        doc = DocxDocument()
        style = doc.styles['Normal']
        style.font.name = '맑은 고딕'
        style.font.size = Pt(11)
        style.paragraph_format.line_spacing = 1.8
        style.paragraph_format.space_after = Pt(4)
        for s in doc.sections:
            s.top_margin = Cm(3); s.bottom_margin = Cm(3)
            s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

        # 표지
        for _ in range(4):
            doc.add_paragraph()
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run(meta.get("name", name))
        r.font.size = Pt(32); r.bold = True
        doc.add_paragraph()
        doc.add_page_break()

        # 챕터 수집 (final > expanded > final_v > draft 순)
        planned = _parse_chapter_plan(_read_text(path / "outline/chapters.md"))
        drafts_dir = path / "drafts"
        final_dir = path / "final"
        total_chars = 0

        for ch in planned:
            cid = ch["id"]
            # 파일 우선순위
            candidates = []
            if final_dir.exists():
                candidates.append(final_dir / f"{cid}.md")
            if drafts_dir.exists():
                for suffix in ["_final.md", "_expanded.md", "_v2.md", "_draft.md"]:
                    candidates.append(drafts_dir / f"{cid}{suffix}")

            content = ""
            for c in candidates:
                if c.exists():
                    content = c.read_text(encoding="utf-8")
                    break

            if not content:
                continue

            # 챕터 제목
            doc.add_paragraph()
            cl = doc.add_paragraph()
            cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cl.add_run(f'Chapter {cid[2:]}')
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

            title_text = ch.get("title") or cid
            ct = doc.add_paragraph()
            ct.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = ct.add_run(title_text)
            r.font.size = Pt(18); r.bold = True
            doc.add_paragraph()

            # 본문
            body_started = False
            for line in content.split('\n'):
                if line.startswith('# '):
                    body_started = True
                    continue
                if not body_started:
                    continue
                stripped = line.strip()
                if not stripped:
                    continue
                if stripped == '* * *':
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.add_run('*   *   *')
                    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                    doc.add_paragraph()
                    continue
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(1)
                r = p.add_run(stripped)
                r.font.size = Pt(11)
                total_chars += len(stripped)

            doc.add_page_break()

        # 끝
        doc.add_paragraph()
        end = doc.add_paragraph()
        end.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = end.add_run('— 끝 —')
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        # 임시 파일에 저장 후 반환
        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        doc.save(tmp.name)
        tmp.close()

        safe_name = name.replace(" ", "_")
        return FileResponse(
            tmp.name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"{safe_name}_전체본.docx",
        )

    # ─────────── 집필 잡 & SSE ───────────
    @app.post("/api/projects/{name}/write")
    async def api_start_write(name: str, request: Request) -> JSONResponse:
        _project_path(name)
        body = await request.json()
        chapter_id = body.get("chapter_id", "").strip()
        if not chapter_id:
            raise HTTPException(status_code=400, detail="chapter_id required")
        job = job_manager.start_write(name, chapter_id)
        return JSONResponse({"job_id": job.id, "status": job.status})

    @app.get("/api/jobs/{job_id}")
    def api_job_snapshot(job_id: str) -> JSONResponse:
        job = job_manager.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="job not found")
        return JSONResponse(job.snapshot())

    @app.get("/api/jobs")
    def api_jobs_recent() -> JSONResponse:
        return JSONResponse({"jobs": job_manager.list_recent()})

    @app.get("/api/jobs/{job_id}/events")
    async def api_job_events(job_id: str) -> StreamingResponse:
        job = job_manager.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="job not found")

        async def event_gen():
            loop = asyncio.get_event_loop()
            gen = job.stream()
            while True:
                item = await loop.run_in_executor(None, _next_or_none, gen)
                if item is _STOP:
                    # 종료 이벤트
                    yield f"event: end\ndata: {json.dumps(job.snapshot(), ensure_ascii=False)}\n\n"
                    return
                if item is None:
                    yield ": keepalive\n\n"  # SSE 코멘트
                    continue
                payload = json.dumps(item.to_dict(), ensure_ascii=False)
                yield f"data: {payload}\n\n"

        return StreamingResponse(
            event_gen(),
            media_type="text/event-stream",
            headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
        )

    return app


_STOP = object()


def _next_or_none(gen):
    try:
        return next(gen)
    except StopIteration:
        return _STOP


def run(host: str = "127.0.0.1", port: int = 8765) -> None:
    import uvicorn
    uvicorn.run(
        "novelcraft.dashboard:create_app",
        host=host,
        port=port,
        factory=True,
        log_level="warning",
    )
