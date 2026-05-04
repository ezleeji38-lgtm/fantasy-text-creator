"""베타리딩 보완본 docx export.

novelcraft.dashboard 의 export 로직을 재사용하여 한 권의 docx 생성.
"""
import json
import re
import sys
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CHAPTER_HEADER_RE = re.compile(r"^##[ \t]+(ch\d{2,3})[ \t]*[—–\-:]?[ \t]*([^\n]*)$", re.MULTILINE)
TITLE_LINE_RE = re.compile(r"^[-*]?\s*\*\*제목\*\*\s*:?\s*(.+)$", re.MULTILINE)


def parse_chapter_plan(text: str) -> list[dict]:
    matches = list(CHAPTER_HEADER_RE.finditer(text))
    chapters = []
    for i, m in enumerate(matches):
        cid = m.group(1)
        same_line_title = m.group(2).strip().lstrip("—–-: ").strip()
        # 다음 챕터 헤더까지 본문에서 "- **제목**: ..." 라인 찾기
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end]
        body_title_match = TITLE_LINE_RE.search(body)
        body_title = body_title_match.group(1).strip() if body_title_match else ""
        title = same_line_title or body_title or cid
        chapters.append({"id": cid, "title": title})
    return chapters


def export(project_dir: Path, out_path: Path, label: str) -> int:
    meta = json.loads((project_dir / "project.json").read_text(encoding="utf-8"))
    chapters_md = (project_dir / "outline/chapters.md").read_text(encoding="utf-8")
    planned = parse_chapter_plan(chapters_md)

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.8
    style.paragraph_format.space_after = Pt(4)
    for s in doc.sections:
        s.top_margin = Cm(3)
        s.bottom_margin = Cm(3)
        s.left_margin = Cm(2.5)
        s.right_margin = Cm(2.5)

    # 표지
    for _ in range(4):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(meta.get("name", project_dir.name))
    r.font.size = Pt(32)
    r.bold = True
    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run(label)
    rs.font.size = Pt(13)
    rs.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    doc.add_page_break()

    drafts_dir = project_dir / "drafts"
    final_dir = project_dir / "final"
    total = 0

    for ch in planned:
        cid = ch["id"]
        candidates = []
        if final_dir.exists():
            candidates.append(final_dir / f"{cid}.md")
        if drafts_dir.exists():
            for suffix in ["_final.md", "_expanded.md", "_v2.md", "_draft.md"]:
                candidates.append(drafts_dir / f"{cid}{suffix}")

        content = ""
        for c in candidates:
            if c.exists():
                content = c.read_text(encoding="utf-8")
                break
        if not content:
            continue

        doc.add_paragraph()
        # 프롤로그(ch00)는 Chapter 라벨 생략
        if cid != "ch00":
            cl = doc.add_paragraph()
            cl.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cl.add_run(f"Chapter {cid[2:]}")
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        ct = doc.add_paragraph()
        ct.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = ct.add_run(ch.get("title") or cid)
        r.font.size = Pt(18)
        r.bold = True
        doc.add_paragraph()

        body_started = False
        for line in content.split("\n"):
            if line.startswith("# "):
                body_started = True
                continue
            if not body_started:
                continue
            stripped = line.strip()
            if not stripped:
                continue
            if stripped == "* * *":
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run("*   *   *")
                r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(1)
            r = p.add_run(stripped)
            r.font.size = Pt(11)
            total += len(stripped)

        doc.add_page_break()

    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end.add_run("— 끝 —")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return total


if __name__ == "__main__":
    project = Path(sys.argv[1])
    out = Path(sys.argv[2])
    label = sys.argv[3] if len(sys.argv) > 3 else ""
    chars = export(project, out, label)
    print(f"saved: {out}  ({chars:,} chars)")
